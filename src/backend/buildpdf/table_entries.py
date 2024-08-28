from docx import Document
from typing import List
from pydantic import BaseModel


def get_table_entries_in_docx(docx_path, current_table_entries) -> List[str]:
    try:
        doc = Document(docx_path)
    except Exception as e:
        return [[f"Error: {e}"]]
    if len(doc.tables) == 0:
        return []
    if len(doc.tables) > 1:
        raise ValueError("Only one table is allowed in the document.")
    table = doc.tables[0]

    table_entries = get_table_entries(table)

    print("current_table_entries", current_table_entries)
    if current_table_entries:
        current_table_entries = {name: _id for name, _id in current_table_entries}
    else:
        current_table_entries = {}

    if table_entries:
        return [
            [entry.name, current_table_entries.get(entry.name)]
            for entry in table_entries
        ]
    return [[]]


def set_cell_text(cell, text):
    """
    Set the text of a cell in a table.
    Use the formatting of the first run. Delete all other runs.
    """
    paragraph = cell.paragraphs[0]
    paragraph.runs[0].text = text

    if len(paragraph.runs) > 1:
        for run in paragraph.runs[1:]:
            run.clear()


class TableEntryData(BaseModel):
    name: str
    page_start: str
    page_end: str


class TableEntry:
    PAGE_START_COL = 3
    PAGE_END_COL = 4
    NAME_COL = 0

    def __init__(self, table, row_num):
        self.table = table
        self.row_num = row_num
        self.get_data()

    def __repr__(self):
        return f"{self.name} {self.page_start} {self.page_end}"

    def get_data(self):
        self.name = self.table.cell(self.row_num, self.NAME_COL).text
        self.page_start = self.table.cell(self.row_num, TableEntry.PAGE_START_COL).text
        self.page_end = self.table.cell(self.row_num, TableEntry.PAGE_END_COL).text

    def set_page_start(self, page_start):
        set_cell_text(
            self.table.cell(self.row_num, TableEntry.PAGE_START_COL), page_start
        )
        self.page_start = page_start

    def set_page_end(self, page_end):
        set_cell_text(self.table.cell(self.row_num, TableEntry.PAGE_END_COL), page_end)
        self.page_end = page_end

    def to_pydantic(self):
        return {
            "name": self.name,
            "page_start": self.page_start,
            "page_end": self.page_end,
        }


def get_table_entries(table):
    SKIPROWS = 2
    num_rows = len(table.rows)
    entries = []
    for i in range(SKIPROWS, num_rows):
        entry = TableEntry(table, i)
        if entry.page_start and entry.page_end:
            entries.append(entry)
    return entries
