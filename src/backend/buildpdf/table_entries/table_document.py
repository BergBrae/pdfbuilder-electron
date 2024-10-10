from pydantic import BaseModel
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
import os
import re


class TableEntry(BaseModel):
    title: str
    page_start: int
    page_end: int
    level: int  # 0-indexed depth of the bookmark


class TableDocument:
    def __init__(
        self,
        docx_path: str,
        page_start_col: int,
        level_delimiter: str = "   ",
        page_end_col: int | None = None,
        skiprows: int = 2,
    ):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.table = self.doc.tables[0]
        self.table_entries: list[TableEntry] = []
        self.set_page_start_col(page_start_col)
        self.set_page_end_col(page_end_col)
        self.level_delimiter = level_delimiter
        self.skiprows = skiprows

    def delete_row(self, row_index: int):
        tbl = self.table._tbl
        tr = self.table.rows[row_index]._tr
        tbl.remove(tr)

    def set_table_entries(self, table_entries: list[TableEntry]):
        self.table_entries = table_entries
        self.adjust_num_rows()
        self.update_cells()

    def adjust_num_rows(self, num_rows=None):
        if num_rows is None:
            num_rows = len(self.table_entries)

        additional_rows = (
            sum(1 for i in range(num_rows - 1) if self.table_entries[i + 1].level == 0)
            if self.table_entries
            else 0
        )
        total_rows_needed = num_rows + additional_rows
        num_rows_to_add = total_rows_needed - len(self.table.rows) + self.skiprows
        if num_rows_to_add > 0:
            for _ in range(num_rows_to_add):
                row = self.table.add_row()
                for cell in row.cells:
                    cell.text = ""  # Clear text in the newly added row
        elif num_rows_to_add < 0:
            for _ in range(-1 * num_rows_to_add):
                self.delete_row(-1)

    def set_page_start_col(self, page_start_col: int):
        self.page_start_col = (
            page_start_col - 1
        )  # user-specified starts at 1, but we need to adjust for 0-indexing

    def set_page_end_col(self, page_end_col: int | None):
        if page_end_col is not None:
            self.page_end_col = page_end_col - 1
        else:
            self.page_end_col = None

    def update_cells(self):
        from docx.shared import Pt  # Import the Pt class for setting font size

        row_idx_to_clear = []
        row_index = self.skiprows
        max_chars_per_line = 50  # Adjust this value based on your document's formatting
        top_level_counter = 1  # Initialize counter for top-level entries

        for i, entry in enumerate(self.table_entries):
            row = self.table.rows[row_index]
            indented_title = self.level_delimiter * entry.level + entry.title

            # Add numbering for top-level bookmarks
            if entry.level == 0:
                indented_title = f"{top_level_counter}. {indented_title}"
                top_level_counter += 1

            wrapped_title = self.wrap_text(indented_title, max_chars_per_line)

            # Set text and font size for the title
            title_cell = row.cells[0]
            title_cell.text = wrapped_title
            for paragraph in title_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)  # Set the font size to 8 points
                    # Bold the text for top-level bookmarks
                    run.font.bold = entry.level == 0

            # Set text and font size for the page start
            page_start_cell = row.cells[self.page_start_col]
            page_start_cell.text = str(entry.page_start)
            for paragraph in page_start_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)  # Set the font size to 8 points

            # Set text and font size for the page end if applicable
            if self.page_end_col is not None:
                page_end_cell = row.cells[self.page_end_col]
                page_end_cell.text = str(entry.page_end)
                for paragraph in page_end_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)  # Set the font size to 8 points

            row_index += 1
            # Add a blank row if the next entry is a level 0
            if i < len(self.table_entries) - 1 and self.table_entries[i + 1].level == 0:
                row_idx_to_clear.append(row_index)
                row_index += 1

        # Clear the text in rows designated to separate top-level entries
        for row_idx in row_idx_to_clear:
            for cell in self.table.rows[row_idx].cells:
                cell.text = ""

    def wrap_text(self, text, max_chars_per_line):
        import re

        # Extract indentation from the beginning of the text
        match = re.match(r"^(\s*)(.*)", text)
        base_indentation = match.group(1)
        text_without_indent = match.group(2)

        additional_indent = (
            "  "  # One-space additional indentation for subsequent lines
        )
        lines = []
        indentation = base_indentation

        first_line = True
        while len(text_without_indent) > max_chars_per_line - len(indentation):
            split_index = text_without_indent.rfind(
                " ", 0, max_chars_per_line - len(indentation)
            )
            if split_index == -1:
                split_index = max_chars_per_line - len(indentation)
            lines.append(indentation + text_without_indent[:split_index])
            text_without_indent = text_without_indent[split_index:].lstrip()

            if first_line:
                # Update indentation for subsequent lines
                indentation = base_indentation + additional_indent
                first_line = False

        lines.append(indentation + text_without_indent)
        return "\n".join(lines)

    def to_pdf(self) -> PdfReader:
        temp_path_docx = "intermediate.docx"
        temp_path_pdf = "intermediate.pdf"
        self.doc.save(temp_path_docx)
        convert(temp_path_docx, temp_path_pdf)
        reader = PdfReader(temp_path_pdf)
        os.remove(temp_path_docx)
        os.remove(temp_path_pdf)
        return reader

    def save(self):
        path = "output.docx"
        self.doc.save(path)
        return path


if __name__ == "__main__":
    # example data
    num_chapters = 30
    num_levels = 3

    docx_path = r"E:\Merit\pdfbuilder\Inventory Sheet Template.docx"

    bookmarks = []
    for chapter in range(1, num_chapters + 1):
        bookmarks.append(
            TableEntry(
                title=f"Chapter {chapter}",
                page_start=(chapter - 1) * 20 + 1,
                page_end=chapter * 20,
                level=0,
            )
        )
        for section in range(1, num_chapters + 1):
            bookmarks.append(
                TableEntry(
                    title=f"Section {chapter}.{section}",
                    page_start=(chapter - 1) * 20 + (section - 1) * 4 + 1,
                    page_end=(chapter - 1) * 20 + section * 4,
                    level=1,
                )
            )

    td = TableDocument(
        docx_path=docx_path,
        page_start_col=2,
        level_delimiter="   ",
        page_end_col=3,
        skiprows=2,
    )
    td.set_table_entries(bookmarks)
    pdf_reader = td.to_pdf()

    writer = PdfWriter()
    for page in pdf_reader.pages:
        writer.add_page(page)

    with open("output.pdf", "wb") as f:
        writer.write(f)
