from docx import Document
from typing import List
from pydantic import BaseModel


def get_table_entries_in_docx(
    docx_path, current_table_entries, page_start_col, page_end_col
) -> List[str]:
    try:
        doc = Document(docx_path)
    except Exception as e:
        return [[f"Error: {e}"]]
    if len(doc.tables) == 0:
        return None
    table = doc.tables[0]

    if not page_start_col:
        return []

    table_entries = get_table_entries(table, page_start_col, page_end_col)

    if current_table_entries and len(current_table_entries[-1]) > 1:
        current_table_entries = {name: _id for name, _id in current_table_entries}
    else:
        current_table_entries = {}

    if table_entries:
        return [
            [entry.name, current_table_entries.get(entry.name)]
            for entry in table_entries
        ]
    return []


def set_cell_text(cell, text):
    """
    Set the text of a cell in a table.
    Use the formatting of the first run. Delete all other runs.
    """
    paragraph = cell.paragraphs[0]
    try:
        paragraph.runs[0].text = str(text) if text else ""
    except IndexError:
        paragraph.add_run(str(text) if text else "")

    if len(paragraph.runs) > 1:
        for run in paragraph.runs[1:]:
            run.clear()


class TableEntryData(BaseModel):
    name: str
    page_start: str
    page_end: str


class TableEntry:
    NAME_COL = 0

    def __init__(self, table, row_num, page_start_col, page_end_col):
        self.table = table
        self.row_num = row_num
        self.page_start_col = page_start_col
        self.page_end_col = page_end_col
        self.get_data()

    def __repr__(self):
        return f"{self.name} {self.page_start} {self.page_end}"

    def get_data(self):
        self.name = self.table.cell(self.row_num, self.NAME_COL).text
        self.page_start = self.table.cell(self.row_num, self.page_start_col).text
        self.page_end = (
            (self.table.cell(self.row_num, self.page_end_col).text)
            if self.page_end_col
            else ""
        )

    def set_page_start(self, page_start):
        set_cell_text(self.table.cell(self.row_num, self.page_start_col), page_start)
        self.page_start = page_start

    def set_page_end(self, page_end):
        if page_end:
            set_cell_text(self.table.cell(self.row_num, self.page_end_col), page_end)
        self.page_end = page_end

    def to_pydantic(self):
        return {
            "name": self.name,
            "page_start": self.page_start,
            "page_end": self.page_end,
        }


def get_table_entries(table, page_start_col, page_end_col):
    SKIPROWS = 2
    num_rows = len(table.rows)
    entries = []
    for i in range(SKIPROWS, num_rows):
        entry = TableEntry(table, i, page_start_col, page_end_col)
        if entry.page_start:
            entries.append(entry)
    return entries
