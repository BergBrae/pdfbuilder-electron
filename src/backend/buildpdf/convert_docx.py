# convert_docx.py
import os
from docx import Document
from python_docx_replace import docx_replace, docx_get_keys
from docx2pdf import convert
from PyPDF2 import PdfReader
import shutil

# from buildpdf.table_entries.table_entries import TableEntry, TableEntryData
from buildpdf.table_entries.table_document import TableDocument, TableEntry
from schema import BookmarkItem


def get_variables_in_docx(docx_path):
    # Create an object of the Document class
    document = Document(docx_path)

    # Get the keys in the DOCX file
    try:
        keys = docx_get_keys(document)
    except Exception as e:
        keys = [f"Error: {e}"]
    return keys


def replace_text_in_docx(docx_path, replacements, output_path=None):
    # Create an object of the Document class
    document = Document(docx_path)

    # Replace the text in the DOCX file
    docx_replace(document, **replacements)

    # Save the modified DOCX file
    if output_path:
        modified_docx_path = output_path
    else:
        modified_docx_path = docx_path.replace(".docx", "_modified.docx")
    document.save(modified_docx_path)
    return modified_docx_path


def convert_docx_to_pdf(docx_path):
    """Converts a DOCX file to PDF and returns the path to the PDF.
    Ensures the PDF file is saved with a .pdf extension.
    """
    # Ensure the output path has a .pdf extension
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    print(f"Converting {docx_path} to {pdf_path}")
    try:
        convert(docx_path, pdf_path)
        if not os.path.exists(pdf_path):
            print(
                f"Warning: PDF conversion seemed successful but file not found at {pdf_path}"
            )
            # Try the default path provided by docx2pdf if our path failed
            default_pdf_path = docx_path.replace(".docx", ".pdf")
            if os.path.exists(default_pdf_path):
                print(
                    f"Found PDF at default location: {default_pdf_path}. Moving to target: {pdf_path}"
                )
                shutil.move(default_pdf_path, pdf_path)
            else:
                # Re-raise or handle the error appropriately if file truly doesn't exist
                raise FileNotFoundError(
                    f"PDF file not found after conversion: {pdf_path}"
                )

        return pdf_path
    except Exception as e:
        print(f"Error during DOCX to PDF conversion for {docx_path}: {e}")
        # Consider if we should try to return a partially created path or None
        # If the error means no PDF was created, returning None might be best
        # If conversion library creates pdf at different path, try to find it
        alt_path = docx_path.replace(".docx", ".pdf")
        if os.path.exists(alt_path):
            print(
                f"Conversion failed but found PDF at alternative path: {alt_path}. Moving."
            )
            try:
                shutil.move(alt_path, pdf_path)
                return pdf_path
            except Exception as move_error:
                print(f"Failed to move alternative PDF: {move_error}")

        raise  # Re-raise the original exception if we can't recover


def update_table_of_contents(
    docx_path, table_entries, page_start_col, page_end_col, total_pages
):
    doc = Document(docx_path)
    table = doc.tables[0]
    num_rows = len(table.rows)
    for i in range(num_rows):
        entry = TableEntry(table, i, page_start_col, page_end_col)
        if entry.name in table_entries:
            page_start = table_entries[entry.name][0]
            page_end = table_entries[entry.name][1]
            if page_end == -1:
                page_end = total_pages
            entry.set_page_start(page_start)
            entry.set_page_end(page_end)
    new_docx_path = docx_path.replace(".docx", "_updated_toc.docx")
    doc.save(new_docx_path)
    return new_docx_path


def convert_bookmark_data_to_table_entries(
    bookmark_data: list[BookmarkItem],
) -> list[TableEntry]:
    def get_level(bookmark: BookmarkItem, level: int = 0) -> int:
        if bookmark.parent is None:
            return level
        return get_level(bookmark.parent, level + 1)

    table_entries = []
    for bookmark in bookmark_data:
        if bookmark.include_in_table_of_contents:
            level = get_level(bookmark)
            table_entry = TableEntry(
                title=bookmark.title,
                page_start=bookmark.page,
                page_end=(
                    bookmark.page_end
                    if bookmark.page_end is not None
                    else bookmark.page
                ),
                level=level,
            )
            table_entries.append(table_entry)
    return table_entries


def convert_docx_template_to_pdf(
    docx_path,
    replacements=None,
    page_start_col=None,
    page_end_col=None,
    page_number_offset=0,
    total_pages=None,
    is_table_of_contents=False,
    bookmark_data=None,
    save_modified_to=None,  # Path to save the modified docx
):
    """Processes a DOCX template: applies replacements, updates TOC if needed, converts to PDF.

    Args:
        docx_path (str): Path to the original DOCX template.
        replacements (dict, optional): Key-value pairs for text replacement.
        save_modified_to (str, optional): If provided, save the modified DOCX to this path.
                                        Otherwise, an intermediate file is used and deleted.
        is_table_of_contents (bool): Flag indicating if the DOCX is a table of contents.
        bookmark_data (list, optional): Data needed to update the table of contents.
        page_start_col, page_end_col, page_number_offset, total_pages: TOC related args.

    Returns:
        tuple: (pdf_reader, num_pages, created_pdf_path, created_docx_path)
               Returns paths to the successfully created files (PDF and modified DOCX).
               Returns None for paths if creation failed or wasn't applicable.
    """
    intermediate_files = []
    current_docx_path = docx_path
    created_docx_path = None
    created_pdf_path = None

    # --- Step 1: Handle Replacements ---
    if replacements:
        try:
            # Determine the output path for the modified DOCX
            if save_modified_to:
                modified_docx_path = save_modified_to
                created_docx_path = save_modified_to  # This is the final intended path
            else:
                # Create a temporary path for the intermediate modified file
                temp_modified_path = current_docx_path.replace(
                    ".docx", "_modified_temp.docx"
                )
                modified_docx_path = temp_modified_path
                intermediate_files.append(temp_modified_path)

            # Perform replacements and save to the determined path
            replace_text_in_docx(current_docx_path, replacements, modified_docx_path)

            # Update the path to use for the next steps
            current_docx_path = modified_docx_path

            # Verify creation only if it's the final intended path
            if save_modified_to and not os.path.exists(created_docx_path):
                print(
                    f"Warning: Modified DOCX specified path {created_docx_path} not found after replacement."
                )
                created_docx_path = None  # Mark as not created

        except Exception as e:
            print(f"Error replacing text in DOCX '{docx_path}': {str(e)}")
            # If replacements fail, we might want to stop or proceed with original?
            # For now, proceed with the current_docx_path (which might be original)
            # but ensure created_docx_path is None if save_modified_to was provided
            if save_modified_to:
                created_docx_path = None

    # --- Step 2: Handle Table of Contents ---
    if is_table_of_contents:
        try:
            # Use the current state of the DOCX (original or modified)
            table_doc = TableDocument(
                docx_path=current_docx_path,
                page_start_col=page_start_col,
                page_end_col=page_end_col,
                skiprows=2,
                page_number_offset=page_number_offset,
            )
            if bookmark_data:
                table_entries = convert_bookmark_data_to_table_entries(bookmark_data)
                table_doc.set_table_entries(table_entries)
                table_doc.adjust_num_rows()

            # Save the TOC updates.
            # If we already have a final modified path, save there.
            # Otherwise, create a new intermediate file.
            if created_docx_path:
                toc_updated_path = created_docx_path
            else:
                toc_updated_path = current_docx_path.replace(
                    ".docx", "_toc_updated_temp.docx"
                )
                if toc_updated_path not in intermediate_files:
                    intermediate_files.append(toc_updated_path)

            table_doc.save(output_path=toc_updated_path)
            current_docx_path = toc_updated_path  # Update path for PDF conversion

            # If we intended to save the final docx, verify TOC update saved correctly
            if save_modified_to and not os.path.exists(created_docx_path):
                print(
                    f"Warning: Modified DOCX with TOC specified path {created_docx_path} not found."
                )
                created_docx_path = None

        except Exception as e:
            print(
                f"Error updating table of contents for '{current_docx_path}': {str(e)}"
            )
            # If TOC update fails, proceed with the current DOCX for PDF conversion
            # If we intended to save the final docx, mark as failed
            if save_modified_to:
                created_docx_path = None

    # --- Step 3: Convert to PDF ---
    pdf_reader = None
    num_pages = 0
    try:
        # Convert the final state of the DOCX to PDF
        created_pdf_path = convert_docx_to_pdf(current_docx_path)

        # Read the generated PDF
        if created_pdf_path and os.path.exists(created_pdf_path):
            pdf_reader = PdfReader(created_pdf_path)
            num_pages = len(pdf_reader.pages)
            print(f"Successfully created and read PDF: {created_pdf_path}")
        else:
            print(
                f"PDF conversion failed or file not found for DOCX: {current_docx_path}"
            )
            created_pdf_path = None  # Ensure path is None if creation failed

    except Exception as e:
        print(f"Error converting DOCX '{current_docx_path}' to PDF: {str(e)}")
        created_pdf_path = None  # Mark PDF as not created on error

    # --- Step 4: Clean up Intermediate Files ---
    for file_path in intermediate_files:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                print(f"Cleaned up intermediate file: {file_path}")
            except Exception as e:
                print(f"Error removing intermediate file {file_path}: {str(e)}")

    # --- Step 5: Return Results ---
    # Ensure created_docx_path is None if the file doesn't actually exist
    if created_docx_path and not os.path.exists(created_docx_path):
        print(
            f"Final check: Modified DOCX path {created_docx_path} does not exist. Setting path to None."
        )
        created_docx_path = None

    # Return the PDF reader, page count, and paths to the final created files
    return (pdf_reader, num_pages, created_pdf_path, created_docx_path)
